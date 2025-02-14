from docx import Document
from docx.shared import Pt, Inches

from excel_test import *

import pandas as pd 

def change_text(doc, position, message, font="Arial", size=10):
    # Access the specific paragraph you want to replace (e.g., the 3rd paragraph)
    paragraph = doc.paragraphs[position]

    # Clear any existing runs in the paragraph
    paragraph.clear()  # Clears the content and runs in the paragraph

    # Add a new run with the desired text
    new_run = paragraph.add_run(message)

    # Set font and size for the new run
    new_run.font.name = font  # Set font to Arial
    new_run.font.size = Pt(size)   # Set font size to 12 points


def create_cap_call_pdf(fund_info, output_name):
    # Load the document
    doc = Document('cap_call_template.docx')

    wire_instructions = {
        "Bank Name:" : "Chase",
        "Bank Address:" : "277 Park Ave New York, NY 10172",
        "ABA Number:" : "021000021 (Domestic Wires)",
        "Account Name:" : fund_info["Fund Name"],
        "Account Number:" : "932859662",
        "SWIFT Code:" : "CHASUS33 (International Wires)"
    }


    for section in doc.sections:
        header = section.header
        
        header_paragraph = header.paragraphs[0]
        header_paragraph.clear()

        ex = r"C:\Users\ppark\OneDrive - GP Fund Solutions, LLC\Desktop\button.png"
        ex2 = r'C:\Users\ppark\OneDrive - GP Fund Solutions, LLC\Desktop\gpes-file-engine-crms\aea-logo.png'
        header_paragraph.add_run().add_picture(ex2, width=Inches(1.5))  # Adjust the size as needed
    
    """
    # Inspect paragraphs and runs
    for i, para in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: {para.text}")
        for j, run in enumerate(para.runs):
            print(f"  Run {j}: {run.text}")
    """

    """
    # Inspect tables
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            for cell in row.cells:
                print(f"  Cell: {cell.text}")
    """


    #change header
    header = doc.tables[0]

    header.rows[2].cells[1].text = fund_info["Fund Name"]
    header.rows[4].cells[1].text = fund_info["Re"]
    header.rows[6].cells[1].text = fund_info["Notice Date"]

    header.rows[8].cells[1].text = ""
    header_fund_info = header.rows[8].cells[1].paragraphs[0].add_run(fund_info["Due Date"])
    header_fund_info.bold = True
    header_fund_info.font.name = "Arial"
    header_fund_info.font.size = Pt(10)

    """
    index = 0
    for row in header.rows:
        for cell in row.cells:
            if index % 2 == 0:
                index += 1
                continue
            if cell.text == "":
                index += 1
                continue

            cell.text = "Test"

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'  # Set the font
                    run.font.size = Pt(10)   # Set the font size to 12 points
            index += 1
    """

    message= f"""In accordance with Section 3.1 of the Amended and Restated Limited Partnership Agreement of the Fund dates {fund_info["Notice Date"]} (the "Agreement"), the Partnership is calling capital for Investments and Management Fees, and Partnership Expenses. Capitalized terms used but not defined in this notice are defined in the Agreement."""
    change_text(doc, 3, message)


    instructions = doc.tables[1]

    """
    # Start iterating from the second row (index 1 instead of 0)
    for index, row in enumerate(instructions.rows):
        # Skip the first row (index 0)
        if index == 0:
            continue
        
        i = 0
        # Iterate through the cells in the row
        for cell in row.cells:
            if i == 0:
                i += 1 
                continue
            # Replace the text in the cell
            cell.text = "test2"

            # Modify the font and size of the text in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"  # Set font to Arial
                    run.font.size = Pt(10)   # Set font size to 10 points
            i += 1
    """

    #Fill in wire instructions
    for row in instructions.rows:
        if row.cells[0].text.strip():
            if str(row.cells[0].text.strip()) in wire_instructions:
                row.cells[1].text = wire_instructions[str(row.cells[0].text.strip())]


    change_text(doc, 15, fund_info["Fund Name"])


    #page 2
    re = fund_info["Re"].split()
    re = ' '.join(re[:3])
    change_text(doc, 21, "Re: " + re)

    #6 columns
    #change 3,4 if 1 is not empty
    #Later, just check if [0] is aligned with a number in a hashmap
    table = doc.tables[2]
    
    for row in table.rows:
        if row.cells[0].text.strip():
            print(row.cells[0].text.strip())
            
            if str(row.cells[0].text.strip()) in fund_info:
                row.cells[2].text = "{:,}".format(fund_info[str(row.cells[0].text.strip())])
        
    

    #go through each investors and create pdfs for them
    print("-----------------------------")
    allocation = pd.read_excel("New Client Allocation Template.xlsx", sheet_name = "Allocation", skiprows=1)
    index = 5
    while (str(allocation.at[index, "Partner Name"]) != "nan"):
        #Gather info
        inv_info = dict()

        inv_info["Investor Name"] = str(allocation.at[index, "Partner Name"])
        header.rows[0].cells[1].text = ""
        header_inv_name = header.rows[0].cells[1].paragraphs[0].add_run(inv_info["Investor Name"])
        header_inv_name.bold = True




        change_text(doc, 20, "Investor: " + inv_info["Investor Name"])

        #investments = int(allocation.at[index, "Investment #1"])
        inv_info["Investment"] = int(allocation.at[index, "Investment #1"])

        #mgmt_fees = int(allocation.at[index, "Gross Mgmt Fee"])
        inv_info["Management Fees"] = int(allocation.at[index, "Gross Mgmt Fee"])

        #pshp_exp = int(allocation.at[index, "Pshp Exp"])
        inv_info["Partnership Expenses"] = int(allocation.at[index, "Pshp Exp"])

        total_amnt = int(allocation.at[index, "Net Amount Due / (Payable)"])
        inv_info["Total Amount Due"] = int(allocation.at[index, "Net Amount Due / (Payable)"])

        cap_commit = int(allocation.at[index, "Commitment"])
        inv_info["Capital Commitment"] = int(allocation.at[index, "Commitment"])

        cum_cap_contributions = int(allocation.at[index, "LTD Ending Contributions"])
        inv_info["Cumulative Capital Contributions"] = int(allocation.at[index, "LTD Ending Contributions"])

        rem_cap_commit = int(allocation.at[index, "Ending Remaining  Commitment"])
        inv_info["Remaining Capital Commitment"] = int(allocation.at[index, "Ending Remaining  Commitment"])


        commit_subj_mgmt_fee = int(allocation.at[index, "LP Commitment"])
        inv_info["Commitment subject to Management Fee"] = int(allocation.at[index, "LP Commitment"])

        mgmt_fee = int(allocation.at[index, "Gross Mgmt Fee"])
        inv_info["Management Fee (1/1/2022 - 3/31/2022)"] = int(allocation.at[index, "Gross Mgmt Fee"])

        mgmt_fee_reduct = int(allocation.at[index, "Mgmt Fee - Offsets"])
        inv_info["Management Fee Reduction"] = int(allocation.at[index, "Mgmt Fee - Offsets"])

        total_mgmt_fee = int(allocation.at[index, "Total Mgmt Fee"])
        inv_info["Total Management Fee, net"] = int(allocation.at[index, "Total Mgmt Fee"])


        #Fill in the info
        instructions.rows[7].cells[1].text = inv_info["Investor Name"]
        message = f"""Your portion of the call is ${inv_info["Total Amount Due"]:,} and is due on {fund_info["Due Date"]}.  Please send your payment by wire transfer in accordance with the instructions provided below."""
        change_text(doc, 7, message)

        for row in table.rows:
            if row.cells[0].text.strip():
                print(row.cells[0].text.strip())
                
                if str(row.cells[0].text.strip()) in inv_info:
                    row.cells[5].text = "{:,}".format(inv_info[str(row.cells[0].text.strip())])



        #print(f"{investor_name} {investments} {mgmt_fees} {pshp_exp} {total_amnt} / {cap_commit} {cum_cap_contributions} {rem_cap_commit} / {commit_subj_mgmt_fee} {mgmt_fee} {mgmt_fee_reduct} {total_mgmt_fee}")

        

        # Save the modified document
        doc.save(output_name + str(index) + ".docx")

        index += 1

    



if __name__ == "__main__":
    fund_info = parse_excel()
    
    create_cap_call_pdf(fund_info, "example_filled")

    
    
